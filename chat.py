import streamlit as st
from dotenv import load_dotenv
import os
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document # type: ignore
from collections import defaultdict
import asyncio
from concurrent.futures import ThreadPoolExecutor

load_dotenv()

genai.configure(api_key=os.getenv("GOOGLE-API-KEY"))

user_interactions = defaultdict(int)

keywords_books = ["Title", "Author", "Chapter", "Summary", "Introduction", "Conclusion", "Page Number", "References", "Abstract", "Publication Date"]
keywords_academic = ["Title", "Abstract", "Introduction", "Methodology", "Results", "Discussion", "Conclusion", "References", "Figures", "Tables", "Keywords", "Author"]
keywords_invoices = ["Invoice Number", "Date", "Total Amount", "Due Date", "Billing Address", "Shipping Address", "Item Description", "Quantity", "Unit Price", "Subtotal", "Tax", "Discount"]
keywords_business = ["Executive Summary", "Objectives", "Introduction", "Scope", "Findings", "Recommendations", "Action Plan", "Appendix", "Budget", "Timeline", "Contacts"]
keywords_general = ["Summary", "Introduction", "Key Points", "Conclusion", "Date", "Author", "Contact Information", "Action Items", "Notes", "Appendix"]

keywords = keywords_books 

executor = ThreadPoolExecutor()

def run_Document():
    st.title("Chat with Documents")
    st.write("Upload multiple documents and start a chat about their content.")
    
    doc_type = st.selectbox(
        "Select the document type:",
        ["Books", "Academic Papers", "Invoices", "Business Documents", "General Documents"]
    )

    if doc_type == "Books":
        keywords = keywords_books
    elif doc_type == "Academic Papers":
        keywords = keywords_academic
    elif doc_type == "Invoices":
        keywords = keywords_invoices
    elif doc_type == "Business Documents":
        keywords = keywords_business
    elif doc_type == "General Documents":
        keywords = keywords_general

    uploaded_files = st.file_uploader("Choose documents...", type=["pdf", "docx"], accept_multiple_files=True)

    if uploaded_files:
        extracted_text = ""
        for uploaded_file in uploaded_files:
            st.write(f"Uploaded file: {uploaded_file.name}")
            file_text = asyncio.run(extract_relevant_sections_async(uploaded_file, keywords))
            if file_text:
                extracted_text += file_text + "\n"
            else:
                st.error(f"Failed to extract text from {uploaded_file.name}.")
        
        if extracted_text:
            input_prompt = st.text_input("Ask a question or generate a quiz from the documents:")

            quiz_generation = st.checkbox("Generate Quiz from Document")

            if st.button("Submit"):
                if input_prompt and not quiz_generation:
                    update_keyword_priorities(input_prompt)
                    response = get_gemini_response(input_prompt, extracted_text)
                    st.write("Response from the AI:")
                    st.write(response)

                if quiz_generation:
                    quiz = generate_quiz_from_document(extracted_text)
                    st.subheader("Quiz generated from the document:")
                    if quiz:
                        for i, qa in enumerate(quiz):
                            st.write(f"Q{i+1}: {qa['question']}")
                            st.write(f"A{i+1}: {qa['answer']}")
                    else:
                        st.write("No quiz could be generated.")

async def extract_relevant_sections_async(uploaded_file, keywords):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(executor, extract_relevant_sections, uploaded_file, keywords)

def extract_relevant_sections(uploaded_file, keywords):
    text = ""
    if uploaded_file.name.endswith(".pdf"):
        text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        text = extract_text_from_docx(uploaded_file)
    return text.strip()

def extract_text_from_pdf(uploaded_file):
    text = ""
    try:
        reader = PdfReader(uploaded_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(uploaded_file):
    text = ""
    try:
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
    return text

def update_keyword_priorities(query):
    for keyword in keywords:
        if keyword.lower() in query.lower():
            user_interactions[keyword] += 1  

def get_gemini_response(prompt, document_text):
    model = genai.GenerativeModel("gemini-pro")
    chat = model.start_chat()

    response = chat.send_message(f"Document: {document_text}\n\nQuestion: {prompt}")

    return response.text



def generate_quiz_from_document(document_text):
    prompt = f"Based on the following text, generate a quiz with multiple questions and detailed answers:\n\n{document_text}"
    model = genai.GenerativeModel("gemini-pro")
    chat = model.start_chat()

    
    response = chat.send_message(prompt)

    
    st.write("Raw AI Response:")
    st.write(response.text)  

    
    quiz = []

    
    response_lines = response.text.split('\n')
    
    current_question = ""
    current_answer = ""
    
    for line in response_lines:
        if line.startswith("Q:"):  
            if current_question:  
                quiz.append({"question": current_question, "answer": current_answer})
            current_question = line.replace("Q:", "").strip()
            current_answer = ""
        elif line.startswith("A:"):  
            current_answer = line.replace("A:", "").strip()
        else:
            if current_answer:
                current_answer += " " + line.strip()

    
    if current_question:
        quiz.append({"question": current_question, "answer": current_answer})

    
    if quiz:
        return quiz
    else:
        st.write("No quiz could be generated. Please check the AI response.")
        return None

